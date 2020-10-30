import datetime
from .models import Paziente, Fattura
import io
from django.http import FileResponse
from django.shortcuts import get_object_or_404
from .utils import data_italiana
from openpyxl import Workbook, load_workbook
from tempfile import NamedTemporaryFile
from docx import Document
from .models import Paziente

numCell = 'B10'
dataCell = 'B11'
nomeCell = 'D9'
indirizzoCell = 'D10'
CAPCell = 'D11'
ComuneCell = 'D12'
CFCell = 'D13'
PIvaCell = 'D14'
TestoCell = 'A20'
IDBolloCell = 'D29'
NettoCell = 'F36'
InpsCell = 'F35'
NoInpsCell = 'F34'
ImportoCell = 'E20'
PagamentoCell = 'A34'
DataPagamentoCell = 'B34'

def crea_dizionario(p:Paziente):
    d = {
        'xxxCognome':p.cognome,
        'xxxNome':p.nome,
        'xxxData':data_italiana(datetime.date.today()),
    }
    d['xxxCodFisc'] = 'C.F.: '+p.codfisc if p.codfisc else ''
    d['xxxPIva'] = 'P.Iva: '+p.piva if p.piva else ''
    if(p.data_nascita and p.paese_nascita and p.provincia_nascita):
        d['xxxNascita'] = (f'nato/a a {p.paese_nascita} ({p.provincia_nascita})'
        f' il {data_italiana(p.data_nascita)}')
    else:
        d['xxxNascita'] = ''
    if(p.paese and p.cap and p.via and p.provincia and p.civico):
        d['xxxResidenza'] = (f'e residente a {p.paese} ({p.provincia}),'
        f' {p.cap}, in {p.via} {p.civico}')
    else:
        d['xxxResidenza'] = ''

    return d

def sostituisci(d:Document, subs):
    for par in d.paragraphs:
        for key in subs:
            if key in par.text:
                inline = par.runs
                for i in range(len(inline)):
                    if key in inline[i].text:
                        text = inline[i].text.replace(key,subs[key])
                        inline[i].text = text
    b = io.BytesIO()
    d.save(b)
    b.seek(0)
    return b

def genera_fattura(request, pk, per_cliente):
    f = get_object_or_404(Fattura, pk=pk)
    p = f.paziente
    nome_file = str(f.numero) + '-' + str(f.get_month()) + '/' + str(f.get_year()) + ' - ' + str(p.cognome) + '.xlsx'
    w = load_workbook(filename='gestione/static/F.xltx')
    s = w.active
    
    s[numCell] = f.numero
    s[dataCell] = f.data
    s[nomeCell] = p.nome + ' ' + p.cognome
    if(p.codfisc):
        s[CFCell] = 'C.F.: ' + p.codfisc
    if(p.piva):
        s[PIvaCell] = 'P.Iva: ' + p.piva
    if(p.via and p.civico and p.cap and p.paese and p.provincia):
        s[indirizzoCell] = p.via + ' N. ' + p.civico
        s[CAPCell] = p.cap
        s[ComuneCell] = p.paese + ' (' + p.provincia + ')'
    s[TestoCell] = f.testo
    s[NettoCell] = f.valore
    s[ImportoCell] = f.valore
    noInps = float(f.valore)/1.04
    s[NoInpsCell] = noInps
    s[InpsCell] = float(f.valore)-noInps

    if not per_cliente and f.data != f.data_incasso:
        if f.data_incasso:
            s[PagamentoCell] = 'Pagamento: '
            s[DataPagamentoCell] = f.data_incasso
        else:
            s[PagamentoCell] = 'Non ancora pagata'

    b = io.BytesIO()
    w.save(b)
    b.seek(0)
    return FileResponse(b, as_attachment=True, filename=nome_file)

def genera_consenso_informato(request, pk, minorenne):
    p = get_object_or_404(Paziente, pk=pk)
    if minorenne:
        doc = Document('gestione/static/CI_m.docx')
    else:
        doc = Document('gestione/static/CI.docx')
    subs = crea_dizionario(p)
    b = sostituisci(doc,subs)
    return FileResponse(b, as_attachment=True, filename='Consenso informato.docx')

def genera_privacy(request, pk, minorenne):
    p = get_object_or_404(Paziente,pk=pk)
    if minorenne:
        doc = Document('gestione/static/P_m.docx')
    else:
        doc = Document('gestione/static/P.docx')
    subs = crea_dizionario(p)
    b = sostituisci(doc,subs)
    return FileResponse(b, as_attachment=True, filename='Privacy.docx')
    